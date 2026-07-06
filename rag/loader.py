"""
文档加载器
==========
支持 PDF、TXT、Markdown、DOCX 等多种文档格式的加载和文本提取。
"""

import hashlib
from pathlib import Path
from typing import List

# PDF
import pypdf

# DOCX
import docx

# TXT/MD 直接读取


class Document:
    """文档数据结构"""

    def __init__(
        self,
        content: str,
        metadata: dict | None = None,
    ):
        self.page_content = content
        self.metadata = metadata or {}

    def __repr__(self):
        source = self.metadata.get("source", "unknown")
        return f"Document(source={source}, chars={len(self.page_content)})"


class DocumentLoader:
    """统一的文档加载器"""

    SUPPORTED_SUFFIXES = {
        ".pdf": "pdf",
        ".txt": "text",
        ".md": "markdown",
        ".docx": "docx",
    }

    @classmethod
    def load_file(cls, file_path: str | Path) -> List[Document]:
        """
        加载单个文件并返回 Document 列表

        Args:
            file_path: 文件路径

        Returns:
            Document 列表 (一个文件可能被拆成多个 Document, 如 PDF 按页)
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        suffix = file_path.suffix.lower()
        if suffix not in cls.SUPPORTED_SUFFIXES:
            raise ValueError(
                f"不支持的文件类型: {suffix}。支持: {list(cls.SUPPORTED_SUFFIXES.keys())}"
            )

        file_type = cls.SUPPORTED_SUFFIXES[suffix]
        loader_method = getattr(cls, f"_load_{file_type}")
        return loader_method(file_path)

    @classmethod
    def load_directory(cls, dir_path: str | Path) -> List[Document]:
        """加载目录下所有支持的文档"""
        dir_path = Path(dir_path)
        all_docs = []
        for suffix in cls.SUPPORTED_SUFFIXES:
            for file_path in dir_path.glob(f"**/*{suffix}"):
                try:
                    docs = cls.load_file(file_path)
                    all_docs.extend(docs)
                    print(f"[Loader] 已加载: {file_path.name} ({len(docs)} 页/段)")
                except Exception as e:
                    print(f"[Loader] 加载失败 {file_path.name}: {e}")
        return all_docs

    # ---- 各格式加载实现 ----

    @staticmethod
    def _load_text(file_path: Path) -> List[Document]:
        """加载纯文本文件"""
        content = file_path.read_text(encoding="utf-8")
        file_hash = hashlib.md5(content.encode()).hexdigest()[:12]
        return [
            Document(
                content=content,
                metadata={
                    "source": file_path.name,
                    "file_type": "txt",
                    "file_hash": file_hash,
                    "page": 0,
                },
            )
        ]

    @staticmethod
    def _load_markdown(file_path: Path) -> List[Document]:
        """加载 Markdown 文件"""
        content = file_path.read_text(encoding="utf-8")
        file_hash = hashlib.md5(content.encode()).hexdigest()[:12]
        return [
            Document(
                content=content,
                metadata={
                    "source": file_path.name,
                    "file_type": "markdown",
                    "file_hash": file_hash,
                    "page": 0,
                },
            )
        ]

    @staticmethod
    def _load_pdf(file_path: Path) -> List[Document]:
        """加载 PDF 文件，按页拆分"""
        documents = []
        file_hash = hashlib.md5(file_path.read_bytes()).hexdigest()[:12]

        reader = pypdf.PdfReader(str(file_path))
        for page_num, page in enumerate(reader.pages, start=1):
            text = page.extract_text()
            if text and text.strip():
                documents.append(
                    Document(
                        content=text.strip(),
                        metadata={
                            "source": file_path.name,
                            "file_type": "pdf",
                            "file_hash": file_hash,
                            "page": page_num,
                        },
                    )
                )

        # 如果按页没提取到内容，尝试整体提取
        if not documents:
            full_text = ""
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    full_text += text + "\n"
            if full_text.strip():
                documents.append(
                    Document(
                        content=full_text.strip(),
                        metadata={
                            "source": file_path.name,
                            "file_type": "pdf",
                            "file_hash": file_hash,
                            "page": 1,
                        },
                    )
                )

        return documents

    @staticmethod
    def _load_docx(file_path: Path) -> List[Document]:
        """加载 DOCX 文件"""
        doc = docx.Document(str(file_path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        content = "\n".join(paragraphs)
        file_hash = hashlib.md5(content.encode()).hexdigest()[:12]
        return [
            Document(
                content=content,
                metadata={
                    "source": file_path.name,
                    "file_type": "docx",
                    "file_hash": file_hash,
                    "page": 0,
                },
            )
        ]
